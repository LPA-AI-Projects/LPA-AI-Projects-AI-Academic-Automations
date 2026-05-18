"""Build .docx bytes for pre/post assessment question banks."""
from __future__ import annotations

from io import BytesIO
from typing import Any


def build_assessment_docx_bytes(
    *,
    course_name: str,
    phase: str,
    difficulty: str,
    questions: list[dict[str, Any]],
) -> bytes:
    from docx import Document

    doc = Document()
    title = f"{course_name.strip() or 'Course'} — {phase.upper()} assessment"
    doc.add_heading(title, 0)
    p = doc.add_paragraph()
    p.add_run("Difficulty: ").bold = True
    p.add_run(difficulty)
    doc.add_paragraph("")

    for i, q in enumerate(questions, start=1):
        if not isinstance(q, dict):
            continue
        stem = str(q.get("question") or "").strip()
        doc.add_heading(f"Question {i}", level=2)
        doc.add_paragraph(stem)
        opts = q.get("options")
        if not isinstance(opts, list):
            opts = []
        for j, opt in enumerate(opts[:4]):
            label = chr(ord("A") + j)
            doc.add_paragraph(f"{label}. {str(opt).strip()}")
        ci = q.get("correct_index", 0)
        try:
            cidx = int(ci)
        except (TypeError, ValueError):
            cidx = 0
        cidx = max(0, min(3, cidx))
        ans = doc.add_paragraph()
        ans.add_run("Correct answer: ").bold = True
        ans.add_run(chr(ord("A") + cidx))
        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
